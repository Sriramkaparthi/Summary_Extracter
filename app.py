import streamlit as st
import pandas as pd
import re
import fitz  # PyMuPDF for PDF processing
import nltk
from nltk.corpus import stopwords
from sklearn.feature_extraction.text import CountVectorizer, TfidfVectorizer
from sklearn.decomposition import LatentDirichletAllocation
from summa.summarizer import summarize
from io import BytesIO
from PIL import Image
import docx
import requests
from bs4 import BeautifulSoup
import numpy as np
from collections import Counter

# Download stopwords
nltk.download('stopwords')
stop_words = set(stopwords.words('english'))

# Load dataset for training
dataset_path = "./Dataset_MAIN.xlsx"
df = pd.read_excel(dataset_path, sheet_name="Sheet1")

# Remove duplicates
df.drop_duplicates(inplace=True)

# Strip whitespace and standardize column names
df.columns = df.columns.str.strip().str.lower()

# Preprocess text function
def preprocess_text(text):
    if isinstance(text, str):
        text = text.lower()
        text = re.sub(r'http\S+', '', text)  # Remove URLs
        text = re.sub(r'[^a-zA-Z\s]', '', text)
        words = text.split()
        words = [word for word in words if word not in stop_words]
        text = ' '.join(words)
    else:
        text = ""
    return text

# Extract keywords from user file using TF-IDF
def extract_keywords(text, top_n=5):
    vectorizer = TfidfVectorizer(stop_words='english', max_features=5000)
    tfidf_matrix = vectorizer.fit_transform([text])
    feature_array = np.array(vectorizer.get_feature_names_out())
    tfidf_sorting = np.argsort(tfidf_matrix.toarray()).flatten()[::-1]
    top_keywords = feature_array[tfidf_sorting][:top_n]
    return list(top_keywords)

# Assign topic dynamically based on extracted keywords
def get_topic_from_keywords(text):
    keywords = extract_keywords(preprocess_text(text))
    return " ".join(keywords) if keywords else "Unknown Topic"

# Extract images from PDF (from main.py)
def extract_images_from_pdf(file_bytes):
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    images = []
    for page in doc:
        for img_index, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            base_image = doc.extract_image(xref)
            img_bytes = base_image["image"]
            img_pil = Image.open(BytesIO(img_bytes))
            if img_pil.width > 100 and img_pil.height > 100:
                images.append(img_pil)
    return images

# Extract images from DOCX (from main.py)
def extract_images_from_docx(file):
    doc = docx.Document(file)
    images = []
    for rel in doc.part.rels:
        if "image" in doc.part.rels[rel].target_ref:
            img_bytes = doc.part.rels[rel].target_part.blob
            img_pil = Image.open(BytesIO(img_bytes))
            if img_pil.width > 100 and img_pil.height > 100:
                images.append(img_pil)
    return images

# Streamlit GUI
st.title("Unsupervised Topic Modeling & Summarization of Scientific Research Documents")
st.write("Upload a document (PDF, DOCX, TXT) or enter a URL below.")

input_type = st.selectbox('Select the document type', ['PDF', 'DOCX', 'Text File', 'Direct Text', 'URL'])

def extract_text_from_pdf(file_bytes):
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    return "\n".join(page.get_text("text") for page in doc)

def extract_text_from_docx(file):
    doc = docx.Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def fetch_text_from_url(url):
    response = requests.get(url)
    soup = BeautifulSoup(response.content, 'html.parser')
    paragraphs = soup.find_all('p')
    return " ".join([para.get_text() for para in paragraphs if para.get_text()])

text = ""
images = []
if input_type == 'PDF':
    uploaded_file = st.file_uploader("Upload your PDF file", type=["pdf"])
    if uploaded_file:
        file_bytes = uploaded_file.read()
        text = extract_text_from_pdf(file_bytes)
        images = extract_images_from_pdf(file_bytes)
elif input_type == 'DOCX':
    uploaded_file = st.file_uploader("Upload your DOCX file", type=["docx"])
    if uploaded_file:
        text = extract_text_from_docx(uploaded_file)
        images = extract_images_from_docx(uploaded_file)
elif input_type == 'Text File':
    uploaded_file = st.file_uploader("Upload your text file", type=["txt"])
    if uploaded_file:
        text = uploaded_file.read().decode('utf-8')
elif input_type == 'Direct Text':
    text = st.text_area("Enter text directly", height=300)
elif input_type == 'URL':
    url = st.text_input("Enter the link to the document (PDF, DOCX, or Webpage)")
    if url:
        text = fetch_text_from_url(url)

if text:
    topic = get_topic_from_keywords(text)
    summary = summarize(text, ratio=0.2)
    
    st.subheader(f"Predicted Topic: {topic}")
    st.subheader("Extracted Text")
    st.text_area("Extracted Text", text, height=300)
    
    st.subheader("Summary")
    st.markdown(f"<div style='text-align: justify; font-size: 16px;'>{summary}</div>", unsafe_allow_html=True)
    
    if images:
        st.subheader("Extracted Images")
        for img in images:
            st.image(img, use_container_width=True)
